"""
    Facial Attribute Analysis Software v0.4
    What's New
        More Performance optimization
        Export Analysis Data in Document and Excel
        Convert Chart Labels to Spanish
    All rights reserved Esteban Morales and Akintola Technologies @ 2024
"""
# IO
import os
import time

# UI
import tkinter as tk
import customtkinter
from tkinter import Frame, RAISED, Button, messagebox, Label, ttk, filedialog, Canvas, LabelFrame

# ML and Image Processing
import tensorflow as tf
import cv2
from PIL import Image, ImageTk
from deepface import DeepFace

# Export
import docx
from docx import Document
import openpyxl
from openpyxl import Workbook

# # Suppress TensorFlow logs
# os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'  # Suppress TensorFlow warnings and info messages
#
# # Disable oneDNN optimizations if causing issues
# os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
#
# # Example of updating deprecated API usage
# loss = tf.compat.v1.losses.sparse_softmax_cross_entropy

# Flags
# flag to indicate of the picture frame is saved
frame_saved = False

# Dictionary for emotion translations
emotion_translations = {
    "neutral": "Neutral",
    "angry": "Enojada",
    "fear": "Miedo",
    "disgust": "Asco",
    "happy": "Feliz",
    "sad": "Triste",
    "surprise": "Sorpresa"
}

# Initialize emotion count dictionary
emotion_counts = {
    "neutral": 0,
    "angry": 0,
    "fear": 0,
    "disgust": 0,
    "happy": 0,
    "sad": 0,
    "surprise": 0
}

current_time = "Nil"


def export_data():
    global current_time, frame_saved

    analytics_data = emotion_counts

    # Calculate total emotions detected and most detected emotion
    total_emotions = sum(analytics_data.values())
    most_detected_emotion = max(analytics_data, key=analytics_data.get)

    # Export docx report
    doc = Document()
    doc.add_heading(f"Facial Attribute Analysis Report")
    doc.add_picture("output/analysed_frame.jpg", width=docx.shared.Inches(4))
    doc.add_paragraph(f"Duration of Facial attribute analysis {current_time}")
    doc.add_paragraph(f"Total emotions detected during analysis: {total_emotions}")
    doc.add_paragraph(f"Most detected emotion during analysis: {most_detected_emotion}")

    # Add a table to the document
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # serial count
    serial = 1

    # Add document header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Serial'
    hdr_cells[1].text = 'Emotion'
    hdr_cells[2].text = 'Count'

    # Create a new Excel workbook and select the active worksheet
    wb = Workbook()
    ws = wb.active

    # Add the table headers
    ws.append(["Facial Attribute Analysis Data"])
    ws.append(["Serial", "Emotion", "Count"])

    # Add data rows
    for emotion, count in analytics_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(serial)
        row_cells[1].text = emotion
        row_cells[2].text = str(count)
        ws.append([str(serial), emotion, str(count)])

        serial = serial + 1

    # save the document
    doc.save("output/analytics-data.docx")

    # Save the sheet
    wb.save("output/analytics-data.xlsx")

    frame_saved = False
    messagebox.showinfo("Export Data", "Analysis exported in Docx and Xslx!")


# functional
def analyze_video(source):
    global emotion_counts, current_time, frame_saved

    # start live time
    start_time = time.time()

    # Get the absolute path to the Haar cascade file to be used during packaging
    current_dir = os.path.dirname(os.path.abspath(__file__))
    cascade_file = os.path.join(current_dir, 'haarcascade_frontalface_default.xml')

    # Load the face cascade classifier
    face_cascade = cv2.CascadeClassifier(cascade_file)

    if source == "webcam":
        # Open the webcam
        cap = cv2.VideoCapture(0)
    else:
        cap = cv2.VideoCapture(source)
    # Check if the webcam is opened successfully
    if not cap.isOpened():
        messagebox.showerror("Error", "Failed to access webcam")
        return

    # Get the frame width and height
    frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    # Define the codec and create VideoWriter object
    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter('output/analysed_vid.mp4', fourcc, 10.0, (frame_width, frame_height))

    while True:
        # Read a frame from the webcam
        ret, frame = cap.read()

        # Check if the frame is read successfully
        if not ret:
            break

        # Convert the frame to grayscale
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

        # Detect faces in the frame
        faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5)

        # Check if any faces are detected
        if len(faces) == 0:
            # print("No faces detected in the frame.")
            continue

        else:
            # Iterate over detected faces
            for x, y, w, h in faces:
                # Draw rectangle around the face
                cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 26), 3)

                # Extract the face region
                face_region = frame[y:y + h, x:x + w]

                try:
                    # Analyze emotions for the face region
                    analyze = DeepFace.analyze(face_region, actions=['emotion'])

                    # language translation
                    language_reference = language.get()

                    # Check if emotions are detected
                    if 'dominant_emotion' in analyze[0]:
                        # Get the dominant emotion
                        dominant_emotion = analyze[0]['dominant_emotion']
                        # Update emotion count
                        emotion_counts[dominant_emotion] += 1

                        if language_reference == "Spanish":
                            dominant_emotion = emotion_translations.get(dominant_emotion, dominant_emotion)

                            # Display the dominant emotion text
                        cv2.putText(frame, dominant_emotion, (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 26), 3)

                    # Update the chart
                    update_chart()

                    # Draw the bar chart with the translated labels
                    draw_bar_chart(chart_canvas, {emotion_translations.get(k, k): v for k, v in emotion_counts.items()})

                except:
                    pass
                    # messagebox.showerror("Error", "Analysis Error")

        # Write the frame into the output video
        out.write(frame)

        # Save the cropped face image to be exported in docx report
        if not frame_saved:
            # Calculate the center and radius of the face
            cx, cy = x + w // 2, y + h // 2
            r = max(w, h) // 2

            # Crop the image to the passport size around the center of the face
            crop_img = frame[max(0, cy - r):min(frame_height, cy + r), max(0, cx - r):min(frame_width, cx + r)]
            crop_img = cv2.resize(crop_img, (413, 531))

            # Save the cropped image
            cv2.imwrite("output/analysed_frame.jpg", crop_img)
            frame_saved = True

        # Convert the frame to RGB format
        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

        # Convert the frame to an ImageTk object
        img = Image.fromarray(frame)
        img = ImageTk.PhotoImage(image=img)

        # Update the video feed label with the new frame
        video_feed.imgtk = img
        video_feed.configure(image=img, width=318, height=456)

        # Refresh the video feed label
        video_feed.update()

        # Calculate the elapsed time
        elapsed_time = time.time() - start_time
        elapsed_minutes = int(elapsed_time // 60)
        elapsed_seconds = int(elapsed_time % 60)
        current_time = f"{elapsed_minutes:02d}:{elapsed_seconds:02d}"
        total_feed_time.config(text=f"Live Time: {elapsed_minutes:02d}:{elapsed_seconds:02d}")

        # Check for key press
        key = cv2.waitKey(1)
        if key == 27:  # Press 'Esc' key to exit
            break

    # Release everything when done
    cap.release()
    out.release()
    cv2.destroyAllWindows()


def upload_video():
    global frame_saved

    file_path = filedialog.askopenfilename(filetypes=[("Video files", "*.mp4;*.avi")])
    if file_path:
        frame_saved = False
        analyze_video(file_path)


def analyze_webcam():
    global frame_saved
    frame_saved = False
    analyze_video("webcam")


def update_chart():
    # update statistical variables
    total_detected_emotions_variable = sum(emotion_counts.values())
    total_detected_emotions.config(text=f"Total Emotions Detected: {total_detected_emotions_variable}")

    # most detected emotion
    if total_detected_emotions_variable > 0:
        max_emotion_variable = max(emotion_counts, key=emotion_counts.get)
        max_emotion.config(text=f"Most Detected Emotion: {max_emotion_variable}")


def draw_bar_chart(canvas, data, width=660, height=450, padding=40):
    # language display option
    language_reference = language.get()

    canvas.delete("all")  # Clear previous drawings

    max_value = max(data.values())
    num_bars = len(data)
    bar_width = (width - 2 * padding) / num_bars  # Adjust width for padding

    if max_value == 0:
        max_value = 1  # Set a default value to avoid division by zero

    for i, (emotion, count) in enumerate(data.items()):
        bar_height = (count / max_value) * (height - 2 * padding)  # Adjust height for padding
        x0 = i * bar_width + padding
        y0 = height - bar_height - padding
        x1 = x0 + bar_width
        y1 = height - padding

        translated_emotion = emotion_translations.get(emotion, emotion)  # Translate emotion label

        canvas.create_rectangle(x0, y0, x1, y1, fill=pry_color)
        canvas.create_text(x0 + bar_width / 2, y1 + 10, text=translated_emotion, anchor=tk.N)
        canvas.create_text(x0 + bar_width / 2, y0 - 10, text=str(count), anchor=tk.S)


# UI Implementation
root = tk.Tk()
root.title("Facial Attribute Analysis Software v.0.4")
root.iconbitmap("fr.ico")
root.geometry("1080x600")
root.resizable(0, 0)

# config
frame_width = 360
frame_height = 600

# UI variables
pry_color = "#514AF5"
sec_color = "#DDDCF2"
_2nd_bgcolor = "white"

# left frame
left_frame = Frame(root, width=frame_width, height=frame_height, bg="white",
                   relief=RAISED, borderwidth=0)
left_frame.place(x=0, y=0)

video_feed = Label(left_frame, width=44, height=30, bg="#C9C9C9")
video_feed.place(x=20, y=20)

# App buttons
upload_btn = customtkinter.CTkButton(
    master=left_frame,
    width=320,
    height=40,
    text="Upload Video File",
    fg_color=sec_color,
    text_color="black",
    command=upload_video
)

upload_btn.place(x=20, y=490)

save_feed = customtkinter.CTkButton(
    master=left_frame,
    width=320,
    height=40,
    text="Start Live Feed",
    fg_color=pry_color,
    command=analyze_webcam
)

save_feed.place(x=20, y=542)

# right frame
right_frame = Frame(root, width=700, height=frame_height, bg="white",
                    relief=RAISED, borderwidth=0)
right_frame.place(x=380, y=0)

# Create a canvas to draw the bar chart
chart_canvas = Canvas(root, width=660, height=460, bg="white")
chart_canvas.place(x=400, y=20)
# chart_canvas.configure(padx=10, pady=10)

# statistical analysis
# analytics_box = LabelFrame(right_frame, text="Statistical Sumarization")
# analytics_box.place(x=20, y=490)

total_feed_time = Label(
    right_frame,
    text="Live Time: 00:00",
    bg=_2nd_bgcolor,
    font=('Arial bold', 10)
)
total_feed_time.place(x=20, y=490)

total_detected_emotions = Label(
    right_frame,
    font=('Arial bold', 10),
    text="Total Emotions Detected: 0",
    bg=_2nd_bgcolor
)
total_detected_emotions.place(x=200, y=490)

max_emotion = Label(
    right_frame,
    text="Most detected Emotion: Nil",
    bg=_2nd_bgcolor,
    font=('Arial bold', 10),
)
max_emotion.place(x=400, y=490)

# Combobox for language reference
app_info = Label(
    master=right_frame,
    text="Language Preference",
    bg="white"
)
app_info.place(x=19, y=520)

language = tk.StringVar()
language_ref = ttk.Combobox(
    master=right_frame,
    width=12,
    textvariable=language,
    state="readonly"
)

language_ref.place(x=20, y=540)
language_ref["values"] = "English Spanish"
language_ref.current(0)

export_stats_btn = customtkinter.CTkButton(
    master=right_frame,
    width=220,
    height=40,
    text="Export Statistics",
    fg_color="#227447",
    hover_color="#06D664",
    command=export_data
)

export_stats_btn.place(x=460, y=542)

# Draw the bar chart with the generated emotion data
draw_bar_chart(chart_canvas, emotion_counts)

root.mainloop()
